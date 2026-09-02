#!/usr/bin/env python3
"""
Translate DOCX text content to Japanese.

- Translates paragraph text
- Translates table cell text
- Preserves document structure/style as much as possible

Usage:
  python translate_docx_to_japanese.py input.docx output_ja.docx
"""

import argparse
from pathlib import Path

from deep_translator import GoogleTranslator
from docx import Document


def is_probably_japanese(text: str) -> bool:
    for ch in text:
        code = ord(ch)
        if (
            0x3040 <= code <= 0x309F
            or 0x30A0 <= code <= 0x30FF
            or 0x4E00 <= code <= 0x9FFF
            or 0xFF66 <= code <= 0xFF9D
        ):
            return True
    return False


def translate_text(text: str, translator: GoogleTranslator) -> str:
    stripped = text.strip()
    if not stripped:
        return text
    if is_probably_japanese(stripped):
        return text
    return translator.translate(text)


def translate_docx(input_path: Path, output_path: Path, source_lang: str = "auto", target_lang: str = "ja"):
    doc = Document(str(input_path))
    translator = GoogleTranslator(source=source_lang, target=target_lang)

    # Paragraphs
    for para in doc.paragraphs:
        if para.text and para.text.strip():
            translated = translate_text(para.text, translator)
            para.text = translated

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text
                if txt and txt.strip():
                    translated = translate_text(txt, translator)
                    cell.text = translated

    doc.save(str(output_path))


def main():
    parser = argparse.ArgumentParser(description="Translate DOCX text to Japanese")
    parser.add_argument("input_docx", help="Path to input DOCX")
    parser.add_argument("output_docx", help="Path to output translated DOCX")
    parser.add_argument("--source-lang", default="auto")
    parser.add_argument("--target-lang", default="ja")
    args = parser.parse_args()

    input_path = Path(args.input_docx)
    output_path = Path(args.output_docx)

    if not input_path.exists():
        raise SystemExit(f"Input DOCX not found: {input_path}")

    translate_docx(input_path, output_path, source_lang=args.source_lang, target_lang=args.target_lang)
    print(f"[DONE] Wrote translated DOCX: {output_path}")


if __name__ == "__main__":
    main()
