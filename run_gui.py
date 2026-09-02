#!/usr/bin/env python3
"""
GUI runner to select files/folders via Explorer/Finder dialogs,
then run OCR/translation tasks with a user-selectable target language.
"""

from pathlib import Path
import tkinter as tk
from tkinter import filedialog, messagebox, ttk

from deep_translator import GoogleTranslator
from docx import Document
from pypdf import PdfReader, PdfWriter
from pypdf.generic import NameObject, TextStringObject, BooleanObject

from translator_core import (
    DEFAULT_TARGET_LANGUAGE_NAME,
    get_supported_languages,
    is_already_target_script,
    ocr_image,
    translate_text,
)


class TranslatorApp:
    def __init__(self, root: tk.Tk):
        self.root = root
        self.root.title("Translator Utility")
        self.root.geometry("480x400")

        self.languages = get_supported_languages()
        if DEFAULT_TARGET_LANGUAGE_NAME not in self.languages:
            self.languages[DEFAULT_TARGET_LANGUAGE_NAME] = "ja"

        self.language_var = tk.StringVar(value=DEFAULT_TARGET_LANGUAGE_NAME)

        frm = tk.Frame(root, padx=20, pady=20)
        frm.pack(fill="both", expand=True)

        tk.Label(frm, text="Choose how you want to run:", font=("Segoe UI", 11, "bold")).pack(
            pady=(0, 12)
        )

        lang_frame = tk.Frame(frm)
        lang_frame.pack(fill="x", pady=(0, 12))
        tk.Label(lang_frame, text="Translate to:").pack(side="left")
        lang_combo = ttk.Combobox(
            lang_frame,
            textvariable=self.language_var,
            values=sorted(self.languages.keys()),
            state="readonly",
            width=28,
        )
        lang_combo.pack(side="left", padx=(8, 0))

        tk.Button(frm, text="Single PNG (pick file)", width=34, command=self.single_file_mode).pack(
            pady=6
        )
        tk.Button(frm, text="Batch PNGs (pick folders)", width=34, command=self.batch_mode).pack(
            pady=6
        )
        tk.Button(frm, text="Translate DOCX (pick file)", width=34, command=self.docx_mode).pack(
            pady=6
        )
        tk.Button(
            frm, text="Translate PDF form fields (pick file)", width=34, command=self.pdf_mode
        ).pack(pady=6)
        tk.Button(frm, text="Exit", width=34, command=root.destroy).pack(pady=6)

    def target_lang_code(self) -> str:
        return self.languages.get(self.language_var.get(), "ja")

    def single_file_mode(self):
        file_path = filedialog.askopenfilename(
            title="Select a PNG image",
            filetypes=[("PNG Images", "*.png"), ("All Files", "*.*")],
        )
        if not file_path:
            return

        target_lang = self.target_lang_code()
        img = Path(file_path)
        out_dir = img.parent / "translated_output"
        out_dir.mkdir(exist_ok=True)

        try:
            ocr_text = ocr_image(img)
            translated_text = translate_text(ocr_text, target_lang=target_lang)

            ocr_out = out_dir / f"{img.stem}.ocr.txt"
            translated_out = out_dir / f"{img.stem}.{target_lang}.txt"
            ocr_out.write_text(ocr_text, encoding="utf-8")
            translated_out.write_text(translated_text, encoding="utf-8")

            messagebox.showinfo(
                "Done",
                f"Processed: {img.name}\n\nSaved:\n{ocr_out}\n{translated_out}",
            )
        except Exception as e:
            messagebox.showerror("Error", str(e))

    def batch_mode(self):
        in_dir = filedialog.askdirectory(title="Select INPUT folder (PNG files)")
        if not in_dir:
            return

        out_dir = filedialog.askdirectory(title="Select OUTPUT folder")
        if not out_dir:
            return

        target_lang = self.target_lang_code()
        in_path = Path(in_dir)
        out_path = Path(out_dir)
        out_path.mkdir(parents=True, exist_ok=True)

        files = sorted(in_path.glob("*.png"))
        if not files:
            messagebox.showwarning("No files", "No PNG files found in selected input folder.")
            return

        translator = GoogleTranslator(source="auto", target=target_lang)
        ok = 0
        fail = 0

        for img in files:
            try:
                ocr_text = ocr_image(img)
                if ocr_text and not is_already_target_script(ocr_text, target_lang):
                    translated_text = translator.translate(ocr_text)
                else:
                    translated_text = ocr_text

                (out_path / f"{img.stem}.ocr.txt").write_text(ocr_text, encoding="utf-8")
                (out_path / f"{img.stem}.{target_lang}.txt").write_text(
                    translated_text, encoding="utf-8"
                )
                ok += 1
            except Exception:
                fail += 1

        messagebox.showinfo("Batch complete", f"Success: {ok}\nFailed: {fail}\nOutput: {out_path}")

    def docx_mode(self):
        file_path = filedialog.askopenfilename(
            title="Select a DOCX file",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
        )
        if not file_path:
            return

        target_lang = self.target_lang_code()
        src = Path(file_path)
        out = src.with_name(f"{src.stem}.{target_lang}.docx")

        try:
            doc = Document(str(src))
            translator = GoogleTranslator(source="auto", target=target_lang)

            def translate_run(txt: str) -> str:
                if txt and txt.strip() and not is_already_target_script(txt, target_lang):
                    return translator.translate(txt)
                return txt

            for para in doc.paragraphs:
                para.text = translate_run(para.text)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell.text = translate_run(cell.text)

            doc.save(str(out))
            messagebox.showinfo("Done", f"Translated DOCX saved:\n{out}")
        except Exception as e:
            messagebox.showerror("Error", str(e))

    def pdf_mode(self):
        file_path = filedialog.askopenfilename(
            title="Select a PDF file",
            filetypes=[("PDF Documents", "*.pdf"), ("All Files", "*.*")],
        )
        if not file_path:
            return

        target_lang = self.target_lang_code()
        src = Path(file_path)
        out = src.with_name(f"{src.stem}.{target_lang}.pdf")

        try:
            reader = PdfReader(str(src))
            writer = PdfWriter()
            for page in reader.pages:
                writer.add_page(page)

            if "/AcroForm" in reader.trailer["/Root"]:
                writer._root_object.update(
                    {NameObject("/AcroForm"): reader.trailer["/Root"]["/AcroForm"]}
                )
                try:
                    writer._root_object["/AcroForm"].update(
                        {NameObject("/NeedAppearances"): BooleanObject(True)}
                    )
                except Exception:
                    pass

            fields = reader.get_fields() or {}
            translator = GoogleTranslator(source="auto", target=target_lang)
            updates = {}

            for field_name, field_obj in fields.items():
                if field_obj.get("/FT") != "/Tx":
                    continue
                value = str(field_obj.get("/V") or "").strip()
                if not value:
                    continue
                if is_already_target_script(value, target_lang):
                    translated = value
                else:
                    translated = translator.translate(value)
                updates[field_name] = TextStringObject(translated)

            for page in writer.pages:
                writer.update_page_form_field_values(page, updates)

            with open(out, "wb") as f_out:
                writer.write(f_out)

            messagebox.showinfo("Done", f"Translated PDF saved:\n{out}")
        except Exception as e:
            messagebox.showerror("Error", str(e))


def main():
    root = tk.Tk()
    TranslatorApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
